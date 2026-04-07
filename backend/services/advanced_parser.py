"""
Advanced Resume Parser for handling various resume formats and layouts
Supports: PDF, DOCX, Images (PNG, JPG, JPEG)
"""

import os
import pdfplumber
import docx
import pytesseract
import cv2
import numpy as np
from PIL import Image, ImageOps
import re
from typing import Dict, List, Tuple, Optional
import warnings
warnings.filterwarnings('ignore')


class AdvancedResumeParser:
    """Enhanced parser for various resume formats with support for non-linear layouts"""
    
    def __init__(self):
        self.section_keywords = {
            'contact': ['contact', 'email', 'phone', 'address', 'linkedin', 'github'],
            'summary': ['summary', 'professional summary', 'objective', 'profile'],
            'experience': ['experience', 'work experience', 'employment', 'professional experience', 'work history'],
            'education': ['education', 'academic', 'degree', 'certification', 'certifications'],
            'skills': ['skills', 'technical skills', 'competencies', 'core competencies', 'expertise', 'languages'],
            'projects': ['projects', 'portfolio', 'notable projects'],
            'awards': ['awards', 'recognition', 'honors', 'achievements'],
            'publications': ['publications', 'research', 'articles'],
            'volunteer': ['volunteer', 'volunteering'],
        }
        self.tesseract_available = self._ensure_tesseract()

    def _ensure_tesseract(self) -> bool:
        """Attempt to locate the Tesseract binary on Windows if not already configured."""
        if pytesseract.pytesseract.tesseract_cmd and os.path.exists(pytesseract.pytesseract.tesseract_cmd):
            return True

        candidates = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            os.path.join(os.environ.get('LOCALAPPDATA', ''), 'Program Files', 'Tesseract-OCR', 'tesseract.exe'),
            os.path.join(os.environ.get('PROGRAMFILES', ''), 'Tesseract-OCR', 'tesseract.exe'),
            os.path.join(os.environ.get('PROGRAMFILES(X86)', ''), 'Tesseract-OCR', 'tesseract.exe')
        ]

        for candidate in candidates:
            if candidate and os.path.exists(candidate):
                pytesseract.pytesseract.tesseract_cmd = candidate
                return True

        print('Warning: Tesseract binary not found. OCR fallback will be disabled.')
        return False
    
    # ============================================
    # IMAGE PREPROCESSING & OCR
    # ============================================
    
    def preprocess_image(self, image_path: str) -> Image.Image:
        """Preprocess image for better OCR accuracy"""
        img = cv2.imread(image_path)
        
        if img is None:
            raise ValueError(f"Could not read image: {image_path}")
        
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        
        denoised = cv2.fastNlMeansDenoising(enhanced, h=10)
        
        adaptive = cv2.adaptiveThreshold(
            denoised,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            11,
            2
        )
        
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        processed = cv2.morphologyEx(adaptive, cv2.MORPH_CLOSE, kernel, iterations=1)
        processed = cv2.morphologyEx(processed, cv2.MORPH_OPEN, kernel, iterations=1)
        
        return Image.fromarray(processed)
    
    def enhance_image_dpi(self, image_path: str, target_dpi: int = 300) -> Image.Image:
        """Enhance image DPI for better OCR"""
        img = Image.open(image_path)
        
        if img.size[0] < 1200 or img.size[1] < 1200:
            scale = max(1200 / img.size[0], 1200 / img.size[1])
            new_size = (int(img.size[0] * scale), int(img.size[1] * scale))
            img = img.resize(new_size, Image.Resampling.LANCZOS)
        
        return img
    
    def _ocr_image(self, pil_image: Image.Image) -> str:
        """Run OCR with several Tesseract PSM modes and return the best result"""
        if not self.tesseract_available:
            return ""

        configs = [
            '--oem 3 --psm 6',
            '--oem 3 --psm 3',
            '--oem 3 --psm 11',
            '--oem 3 --psm 1'
        ]
        best_text = ""
        best_len = 0

        try:
            pil_image = ImageOps.autocontrast(pil_image)
            pil_image = ImageOps.grayscale(pil_image)
        except Exception:
            pass

        for config in configs:
            try:
                text = pytesseract.image_to_string(pil_image, lang='eng', config=config)
                if len(text.strip()) > best_len:
                    best_text = text
                    best_len = len(text.strip())
            except Exception as e:
                print(f"PSM OCR failed ({config}): {e}")

        return best_text
    
    def extract_text_from_image_ocr(self, image_path: str) -> str:
        """Extract text from image using advanced OCR with preprocessing"""
        try:
            enhanced_img = self.enhance_image_dpi(image_path)
            text = self._ocr_image(enhanced_img)
            if text.strip():
                return text
        except Exception as e:
            print(f"Enhanced DPI OCR failed: {e}")
        
        try:
            preprocessed_img = self.preprocess_image(image_path)
            text = self._ocr_image(preprocessed_img)
            if text.strip():
                return text
        except Exception as e:
            print(f"Preprocessed OCR failed: {e}")
        
        try:
            img = Image.open(image_path)
            text = self._ocr_image(img)
            if text.strip():
                return text
        except Exception as e:
            print(f"Basic OCR failed: {e}")
        
        return ""
    
    # ============================================
    # PDF PARSING
    # ============================================
    
    def _extract_text_from_pdf_page(self, page) -> str:
        """Extract text from a PDF page using pdfplumber and OCR fallback"""
        text = page.extract_text() or ""
        if not text or len(text.strip()) < 50:
            try:
                text = page.extract_text(layout=True) or text
            except Exception:
                pass

        if not text or len(text.strip()) < 50:
            try:
                words = page.extract_words()
                if words:
                    lines = []
                    current_line = []
                    last_top = None
                    for word in words:
                        if last_top is not None and abs(word["top"] - last_top) > 8:
                            lines.append(" ".join(current_line))
                            current_line = []
                        current_line.append(word["text"])
                        last_top = word["top"]
                    if current_line:
                        lines.append(" ".join(current_line))
                    text = "\n".join(lines) or text
            except Exception:
                pass

        ocr_text = ""
        try:
            page_image_obj = page.to_image(resolution=300)
            pil_img = None
            if hasattr(page_image_obj, 'original') and page_image_obj.original is not None:
                pil_img = page_image_obj.original
            elif hasattr(page_image_obj, 'image') and page_image_obj.image is not None:
                pil_img = page_image_obj.image
            elif hasattr(page_image_obj, 'render'):
                pil_img = page_image_obj.render()

            if pil_img is not None:
                pil_img = pil_img.convert('RGB')
                ocr_text = self._ocr_image(pil_img)
        except Exception as e:
            print(f"PDF page OCR fallback failed: {e}")

        if len(ocr_text.strip()) > len(text.strip()):
            return ocr_text
        return text or ocr_text
    
    def parse_pdf_advanced(self, path: str) -> Tuple[str, Dict]:
        """Parse PDF with layout detection"""
        text = ""
        layout_info = {
            'has_tables': False,
            'has_multiple_columns': False,
            'pages_count': 0
        }
        
        try:
            with pdfplumber.open(path) as pdf:
                layout_info['pages_count'] = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = self._extract_text_from_pdf_page(page)

                    try:
                        tables = page.extract_tables() or []
                    except Exception:
                        tables = []

                    if tables:
                        layout_info['has_tables'] = True
                        for table in tables:
                            for row in table:
                                page_text += " ".join([str(cell) or "" for cell in row]) + "\n"
                    
                    text += page_text + "\n"
        except Exception as e:
            print(f"PDF parsing error: {e}")
        
        return text, layout_info
    
    # ============================================
    # DOCX PARSING
    # ============================================
    
    def parse_docx_advanced(self, path: str) -> str:
        """Parse DOCX with formatting preservation"""
        text = ""
        try:
            doc = docx.Document(path)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text += " ".join(row_text) + "\n"
        except Exception as e:
            print(f"DOCX parsing error: {e}")
        
        return text
    
    # ============================================
    # SECTION DETECTION & EXTRACTION
    # ============================================
    
    def detect_sections(self, text: str) -> Dict[str, Tuple[int, int]]:
        """Detect resume sections and their boundaries"""
        lines = text.split('\n')
        sections = {}
        
        current_section = None
        section_start = 0
        char_position = 0
        
        for line_idx, line in enumerate(lines):
            line_lower = line.lower().strip()
            char_position_before = char_position
            char_position += len(line) + 1
            
            # Check if line is a section header
            for section_name, keywords in self.section_keywords.items():
                if any(keyword in line_lower for keyword in keywords):
                    # Save previous section
                    if current_section:
                        sections[current_section] = (section_start, char_position_before)
                    
                    current_section = section_name
                    section_start = char_position
                    break
        
        # Save last section
        if current_section:
            sections[current_section] = (section_start, len(text))
        
        return sections
    
    def extract_section_text(self, text: str, section_name: str) -> str:
        """Extract text for a specific section"""
        sections = self.detect_sections(text)
        
        if section_name not in sections:
            return ""
        
        start, end = sections[section_name]
        return text[start:end]
    
    # ============================================
    # STRUCTURED DATA EXTRACTION
    # ============================================
    
    def extract_contact_info(self, text: str) -> Dict:
        """Extract contact information"""
        contact = {}
        
        # Email
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        if email_match:
            contact['email'] = email_match.group()
        
        # Phone
        phone_match = re.search(r'(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})', text)
        if phone_match:
            contact['phone'] = phone_match.group()
        
        # LinkedIn
        linkedin_match = re.search(r'linkedin\.com/in/[\w\-]+', text, re.IGNORECASE)
        if linkedin_match:
            contact['linkedin'] = linkedin_match.group()
        
        # GitHub
        github_match = re.search(r'github\.com/[\w\-]+', text, re.IGNORECASE)
        if github_match:
            contact['github'] = github_match.group()
        
        # Address (basic extraction)
        address_match = re.search(r'(.+?)(?:Email|email|Phone|phone|LinkedIn|github)', text, re.IGNORECASE)
        if address_match:
            contact['address'] = address_match.group(1).strip()
        
        return contact
    
    def extract_experience_entries(self, text: str) -> List[Dict]:
        """Extract work experience entries with intelligent date pattern matching"""
        experience_section = self.extract_section_text(text, 'experience')
        entries = []
        
        if not experience_section:
            return entries
        
        # Split experience section into individual job entries
        # Common separators between entries: newlines followed by a capital letter (job title)
        entry_blocks = re.split(r'\n(?=[A-Z](?:[a-z]+\s+)?(?:at|,|for|–|-|in|\|))', experience_section)
        
        for block in entry_blocks:
            if len(block.strip()) < 10:
                continue
            
            lines = [l.strip() for l in block.split('\n') if l.strip()]
            
            title = None
            company = None
            period = None
            
            # Try to identify job title, company, and date range
            # Usually: Title / Company / Date on separate lines or inline
            
            # First, search for any date pattern in the block
            date_patterns = [
                r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+)?(\d{4})\s*[-–—to]\s*(?:present|current|ongoing|(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+)?(\d{4}))?',
                r'(\d{1,2})[/\-.](\d{4})\s*[-–—to]\s*(\d{1,2})?[/\-.]?(\d{4})?',
            ]
            
            for pattern in date_patterns:
                date_match = re.search(pattern, block, re.IGNORECASE)
                if date_match:
                    period = date_match.group(0)
                    break
            
            # Assign lines to title/company based on content
            for i, line in enumerate(lines):
                line_lower = line.lower()
                
                # Skip if line is just the period
                if period and line == period:
                    continue
                
                # Check if line looks like a job title (contains job keywords)
                job_keywords = ['developer', 'engineer', 'manager', 'analyst', 'specialist', 'lead', 'director', 'architect', 'officer', 'intern', 'associate', 'consultant', 'coordinator', 'administrator']
                if not title and any(keyword in line_lower for keyword in job_keywords):
                    title = line
                # Check if line looks like company name (usually starts with capital, no job keywords)
                elif not company and not any(keyword in line_lower for keyword in job_keywords) and len(line) < 100:
                    company = line
                # If period already found and we have title/company, we're done
                elif period and title and company:
                    break
            
            # Fallback: if we only have one line and it has both job and company, try to split it
            if not title and not company and len(lines) == 1:
                first_line = lines[0]
                # Try to find "Title at/for Company" pattern
                match = re.match(r'([^,|]+?)\s+(?:at|for|@)\s+(.+?)(?:\s+[-–]|$)', first_line, re.IGNORECASE)
                if match:
                    title = match.group(1).strip()
                    company = match.group(2).strip()
                else:
                    title = first_line  # Just use the whole thing as title
            
            # Add entry if we have at least title or company
            if title or company or period:
                entry = {
                    'title': (title or '').strip(),
                    'company': (company or '').strip(),
                    'period': (period or '').strip()
                }
                entries.append(entry)
        
        return entries
    
    def extract_education_entries(self, text: str) -> List[Dict]:
        """Extract education entries"""
        education_section = self.extract_section_text(text, 'education')
        entries = []
        
        # Pattern for education entries
        patterns = [
            r'(Bachelor|Master|PhD|Associate|B\.?S\.?|M\.?S\.?|M\.?B\.?A\.?|B\.?A\.?|M\.?F\.?A\.?)[.\s]+(in\s+)?([A-Za-z\s]+)(?:\s+from\s+|,\s+|@\s+)([A-Za-z0-9\s\.\-]+?)(?:\s+\(([^)]+)\)|$)',
            r'([A-Za-z0-9\s\.\-]+?)\s+(?:University|College|Institute|School)\s*(?:,|\s+)?\s*(?:\(([^)]+)\))?',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, education_section, re.IGNORECASE)
            for match in matches:
                if len(match.groups()) >= 3:
                    entry = {
                        'degree': match.group(1).strip(),
                        'field': match.group(3).strip() if match.group(3) else '',
                        'institution': match.group(4).strip() if match.group(4) else '',
                        'year': match.group(5).strip() if len(match.groups()) >= 5 else ''
                    }
                    entries.append(entry)
        
        return entries
    
    def extract_skills_list(self, text: str) -> List[str]:
        """Extract skills list"""
        skills_section = self.extract_section_text(text, 'skills')
        
        skills = []
        
        # Common skill delimiters
        delimiters = [',', '•', '·', '|', ';', '\n', '-']
        
        for delimiter in delimiters:
            if delimiter in skills_section:
                potential_skills = skills_section.split(delimiter)
                for skill in potential_skills:
                    skill = skill.strip()
                    if 2 < len(skill) < 50:  # Reasonable skill length
                        skills.append(skill)
                break
        
        # Remove duplicates while preserving order
        seen = set()
        unique_skills = []
        for skill in skills:
            skill_lower = skill.lower()
            if skill_lower not in seen:
                seen.add(skill_lower)
                unique_skills.append(skill)
        
        return unique_skills
    
    # ============================================
    # MAIN EXTRACTION METHOD
    # ============================================
    
    def parse_resume(self, filepath: str) -> Dict:
        """Main method to parse resume and extract structured data"""
        
        raw_text = ""
        metadata = {
            'file_format': None,
            'extraction_method': None,
            'has_sections': False,
            'status': 'success'
        }
        
        try:
            # Determine file type and extract text
            if filepath.lower().endswith('.pdf'):
                raw_text, pdf_info = self.parse_pdf_advanced(filepath)
                metadata['file_format'] = 'PDF'
                metadata['extraction_method'] = 'pdfplumber'
                metadata.update(pdf_info)
            
            elif filepath.lower().endswith('.docx'):
                raw_text = self.parse_docx_advanced(filepath)
                metadata['file_format'] = 'DOCX'
                metadata['extraction_method'] = 'python-docx'
            
            elif filepath.lower().endswith(('.png', '.jpg', '.jpeg')):
                raw_text = self.extract_text_from_image_ocr(filepath)
                metadata['file_format'] = filepath.split('.')[-1].upper()
                metadata['extraction_method'] = 'pytesseract_advanced'
            
            else:
                metadata['status'] = 'error'
                metadata['error'] = 'Unsupported file format'
                return metadata
            
            # Clean text
            raw_text = self._clean_text(raw_text)
            
            if not raw_text.strip():
                metadata['status'] = 'error'
                metadata['error'] = 'No text extracted'
                return metadata
            
            # Extract structured data
            sections = self.detect_sections(raw_text)
            metadata['has_sections'] = len(sections) > 0
            metadata['detected_sections'] = list(sections.keys())
            
            structured_data = {
                'raw_text': raw_text,
                'contact_info': self.extract_contact_info(raw_text),
                'experience': self.extract_experience_entries(raw_text),
                'education': self.extract_education_entries(raw_text),
                'skills': self.extract_skills_list(raw_text),
                'sections_detected': sections,
                'metadata': metadata
            }
            
            return structured_data
        
        except Exception as e:
            return {
                'status': 'error',
                'error': str(e),
                'metadata': metadata
            }
    
    # ============================================
    # UTILITY METHODS
    # ============================================
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text while preserving line structure"""
        # Normalize spaces inside lines but preserve line breaks
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r' *\n+ *', '\n', text)
        # Fix common OCR errors in extracted text
        if text.count('|') > 10:
            text = text.replace('|', 'l')
        return text.strip()
    
    def get_full_text(self, filepath: str) -> str:
        """Get full text from resume (backward compatibility)"""
        result = self.parse_resume(filepath)
        if isinstance(result, dict) and 'raw_text' in result:
            return result['raw_text']
        elif isinstance(result, dict) and 'status' in result and result['status'] == 'error':
            return ""
        return str(result)


# Backward compatibility wrapper functions
def extract_text_advanced(filepath: str) -> str:
    """Extract full text from resume file"""
    parser = AdvancedResumeParser()
    return parser.get_full_text(filepath)


def parse_resume_structured(filepath: str) -> Dict:
    """Parse resume and return structured data"""
    parser = AdvancedResumeParser()
    return parser.parse_resume(filepath)
